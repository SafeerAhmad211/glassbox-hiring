"""PDF and DOCX text extraction with positions.

Produces :class:`~glassbox.parse.layout.TextBlock` objects that
:mod:`glassbox.parse.layout` and :mod:`glassbox.parse.diagnostics` consume. Behind an
optional dependency: ``pip install 'glassbox-hiring[parse]'``.

**Why pdfminer.six and not pypdf.** pypdf is a PDF manipulation library; its text
visitor does not reliably report the horizontal component of the text matrix, so
side-by-side columns collapse to a single x coordinate and column detection silently
fails. pdfminer.six is a layout-analysis library and reports true glyph bounding
boxes. Using the wrong one here does not error -- it quietly reports every two-column
resume as clean, which is the exact failure this module exists to detect.

We still use pypdf, but for a different job: its plain ``extract_text()`` is a faithful
example of a **geometry-blind** extractor, so we capture it as the "naive" reading.
That makes the reading-order comparison a contrast between two real-world extractors
rather than between two of our own functions.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from .layout import TextBlock

__all__ = ["extract_blocks", "extract_docx_blocks", "extract_pdf_blocks"]


def extract_blocks(path: Path | str) -> tuple[list[TextBlock], dict[str, Any]]:
    """Extract positioned text from a PDF or DOCX.

    Args:
        path: Path to the document.

    Returns:
        ``(blocks, flags)``. ``flags`` reports parseability-relevant document features:
        ``has_images``, ``has_tables``, ``n_pages``, and for PDFs ``naive_text`` -- the
        output of a geometry-blind extractor.

    Raises:
        ImportError: If the required optional dependency is missing.
        ValueError: If the file type is unsupported.
        OSError: If the file cannot be read.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf_blocks(path)
    if suffix == ".docx":
        return extract_docx_blocks(path)
    raise ValueError(
        f"Unsupported file type {suffix!r}. Supported: .pdf, .docx. "
        "Convert .doc to .docx first."
    )


def extract_pdf_blocks(path: Path | str) -> tuple[list[TextBlock], dict[str, Any]]:
    """Extract positioned text from a PDF using pdfminer.six.

    Args:
        path: Path to the PDF.

    Returns:
        ``(blocks, flags)``.

    Raises:
        ImportError: If pdfminer.six is not installed.
        OSError: If the file cannot be read.
    """
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTFigure, LTImage, LTTextContainer, LTTextLine
    except ImportError as exc:  # pragma: no cover - exercised only without the extra
        raise ImportError(
            "PDF support requires pdfminer.six. Install with: "
            "pip install 'glassbox-hiring[parse]'"
        ) from exc

    blocks: list[TextBlock] = []
    has_images = False
    counter = 0
    page_count = 0

    for page_number, page_layout in enumerate(extract_pages(str(path))):
        page_count += 1
        for element in page_layout:
            if isinstance(element, (LTImage, LTFigure)):
                has_images = True
                continue
            if not isinstance(element, LTTextContainer):
                continue

            for line in element:
                if not isinstance(line, LTTextLine):
                    continue
                text = line.get_text().strip()
                if not text:
                    continue

                # Font size from the tallest glyph on the line; character heights vary
                # within a line and the maximum is the best single descriptor.
                sizes = [
                    float(char.size)
                    for char in line
                    if hasattr(char, "size")
                ]
                blocks.append(
                    TextBlock(
                        text=text,
                        x=float(line.x0),
                        y=float(line.y0),
                        size=max(sizes) if sizes else 0.0,
                        page=page_number,
                        emit_index=counter,
                        width=float(line.x1 - line.x0),
                    )
                )
                counter += 1

    return blocks, {
        "has_images": has_images,
        "has_tables": False,  # PDFs carry no table semantics; geometry is the signal.
        "n_pages": page_count,
        "naive_text": _naive_pdf_text(path),
    }


def _naive_pdf_text(path: Path | str) -> str | None:
    """Text as a geometry-blind extractor produces it.

    Returns ``None`` if pypdf is unavailable -- an optional refinement, not a
    requirement, so its absence must not break extraction.
    """
    try:
        from pypdf import PdfReader
    except ImportError:  # pragma: no cover
        return None

    try:
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except (OSError, ValueError, KeyError):
        # A malformed PDF that pdfminer handled but pypdf rejects is not a reason to
        # fail the whole extraction; we simply lose the naive comparison.
        return None


def extract_docx_blocks(path: Path | str) -> tuple[list[TextBlock], dict[str, Any]]:
    """Extract text from a DOCX.

    DOCX is a flow format with no absolute coordinates, so synthetic positions are
    assigned: paragraphs descend a nominal page and table cells are laid out by column.
    That is enough for the column and section diagnostics, and table presence is
    reported exactly rather than inferred from geometry.

    Args:
        path: Path to the DOCX.

    Returns:
        ``(blocks, flags)``.

    Raises:
        ImportError: If python-docx is not installed.
        OSError: If the file cannot be read.
    """
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - exercised only without the extra
        raise ImportError(
            "DOCX support requires python-docx. Install with: "
            "pip install 'glassbox-hiring[parse]'"
        ) from exc

    document = docx.Document(str(path))
    blocks: list[TextBlock] = []
    y = 720.0
    counter = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(
                TextBlock(text=text, x=50.0, y=y, size=11.0, page=0, emit_index=counter)
            )
            counter += 1
        y -= 16.0

    for table in document.tables:
        for row in table.rows:
            for column_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    blocks.append(
                        TextBlock(
                            text=text,
                            x=50.0 + column_index * 200.0,
                            y=y,
                            size=11.0,
                            page=0,
                            emit_index=counter,
                        )
                    )
                    counter += 1
            y -= 16.0

    has_images = bool(
        document.inline_shapes
        or any("image" in rel.reltype for rel in document.part.rels.values())
    )

    return blocks, {
        "has_images": has_images,
        "has_tables": bool(document.tables),
        "n_pages": 1,
    }
